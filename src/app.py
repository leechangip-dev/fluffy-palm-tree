"""Flask web application for the translation tool."""

import io
import json
import os
import re as _re
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

from flask import Flask, jsonify, render_template, request, send_file

sys.path.insert(0, str(Path(__file__).parent))
from translator import SUPPORTED_LANGUAGES, Translator
from validator import TranslationValidator

app = Flask(__name__)

_translator: Translator | None = None
_validator: TranslationValidator | None = None


def _get_translator() -> Translator:
    global _translator
    if _translator is None:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise RuntimeError("ANTHROPIC_API_KEY environment variable is not set")
        _translator = Translator(api_key=api_key)
    return _translator


def _get_validator(with_translator: bool = False) -> TranslationValidator:
    global _validator
    if with_translator:
        return TranslationValidator(_get_translator())
    if _validator is None:
        _validator = TranslationValidator()
    return _validator


# ── File text extraction ──────────────────────────────────────────────────────

def _extract_docx(data: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    import fitz
    doc = fitz.open(stream=data, filetype="pdf")
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n".join(pages)


def _extract_xlsx_ser(data: bytes) -> list[dict]:
    """Extract SER entries from an Excel file."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    rows = []
    for ws in wb.worksheets:
        headers = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(c) if c is not None else "" for c in row]
                continue
            if all(c is None for c in row):
                continue
            entry = {headers[j]: (row[j] if j < len(row) else None)
                     for j in range(len(headers))}
            rows.append(entry)
    wb.close()
    return rows


def _extract_text(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext in (".txt", ".md"):
        return data.decode("utf-8", errors="replace")
    if ext == ".docx":
        return _extract_docx(data)
    if ext == ".pdf":
        return _extract_pdf(data)
    return data.decode("utf-8", errors="replace")


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html", languages=SUPPORTED_LANGUAGES)


@app.route("/api/extract-text", methods=["POST"])
def api_extract_text():
    """Accept a binary file upload and return extracted plain text."""
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "file required"}), 400
    ext = Path(f.filename).suffix.lower()
    data = f.read()
    try:
        if ext == ".xlsx":
            rows = _extract_xlsx_ser(data)
            text = json.dumps(rows, ensure_ascii=False, indent=2)
        else:
            text = _extract_text(f.filename, data)
        return jsonify({"text": text, "filename": f.filename})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/translate", methods=["POST"])
def api_translate():
    data = request.get_json(force=True)
    text: str = data.get("text", "").strip()
    target_langs: list[str] = data.get("target_langs", ["en"])
    source_lang: str | None = data.get("source_lang") or None
    context: str | None = data.get("context") or None

    if not text:
        return jsonify({"error": "text is required"}), 400

    try:
        translator = _get_translator()
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 500

    results: dict[str, str] = {}
    errors: dict[str, str] = {}

    def _translate_one(lang: str):
        return lang, translator.translate(text, lang, source_lang, context)

    with ThreadPoolExecutor(max_workers=len(target_langs)) as pool:
        futures = {pool.submit(_translate_one, lang): lang for lang in target_langs}
        for future in as_completed(futures):
            try:
                lang, translated = future.result()
                results[lang] = translated
            except Exception as e:
                errors[futures[future]] = str(e)

    return jsonify({"results": results, "errors": errors})


@app.route("/api/validate", methods=["POST"])
def api_validate():
    data = request.get_json(force=True)
    source_text: str = data.get("source_text", "").strip()
    translated_texts: dict[str, str] = data.get("translated_texts", {})
    check_quality: bool = bool(data.get("check_quality", False))

    if not source_text:
        return jsonify({"error": "source_text is required"}), 400
    if not translated_texts:
        return jsonify({"error": "translated_texts is required"}), 400

    try:
        validator = _get_validator(with_translator=check_quality)
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 500

    results: dict[str, dict] = {}

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)
        src_file = tmp / "source.md"
        src_file.write_text(source_text, encoding="utf-8")

        for lang, translated in translated_texts.items():
            trl_file = tmp / f"source.{lang}.md"
            trl_file.write_text(translated, encoding="utf-8")

            vr = validator.validate_file(
                src_file, lang,
                translated_path=trl_file,
                check_quality=check_quality,
            )

            results[lang] = {
                "complete": vr.is_complete,
                "passed": vr.passed,
                "issues": vr.completeness_issues,
                "quality_score": vr.quality_score,
                "quality_feedback": vr.quality_feedback,
            }

    return jsonify({"results": results})


@app.route("/api/detect", methods=["POST"])
def api_detect():
    data = request.get_json(force=True)
    text: str = data.get("text", "").strip()
    if not text:
        return jsonify({"error": "text is required"}), 400
    try:
        translator = _get_translator()
        lang = translator.detect_language(text)
        return jsonify({"lang": lang, "name": SUPPORTED_LANGUAGES.get(lang, lang)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/patent-verify", methods=["POST"])
def api_patent_verify():
    data = request.get_json(force=True)
    source_text: str = data.get("source_text", "").strip()
    translation_text: str = data.get("translation_text", "").strip()
    notes_text: str = data.get("notes_text", "").strip()
    drawing_text: str = data.get("drawing_text", "").strip()
    ser_data: str = data.get("ser_data", "").strip()

    if not source_text:
        return jsonify({"error": "원문(일본어) 텍스트가 필요합니다."}), 400
    if not translation_text:
        return jsonify({"error": "번역문(영어) 텍스트가 필요합니다."}), 400

    try:
        translator = _get_translator()
    except RuntimeError as e:
        return jsonify({"error": str(e)}), 500

    notes_block   = f"\n\n[번역자 메모]\n{notes_text[:2000]}"   if notes_text   else ""
    drawing_block = f"\n\n[도면 PDF 텍스트]\n{drawing_text[:2000]}" if drawing_text else ""
    ser_block     = f"\n\n[SER 데이터]\n{ser_data[:2000]}"      if ser_data     else ""
    drawing_instr = """
8. Drawing callout check: verify that reference numerals in the JP spec match the EN translation
   and the drawing PDF. Report mismatches in "drawing_mismatches".""" if drawing_text else ""

    prompt = f"""You are a senior patent translation verifier (Japanese→English PCT).

Verify sentence-by-sentence, checking:
1. Strict literal fidelity — no fluency smoothing
2. Patent terminology accuracy and consistency
3. Claims conventions: a/an/the articles, "comprising"/"wherein"/"configured to"
4. Rephrasing, omissions, additions, paraphrasing
5. Grammar / syntax
6. Translator notes (if provided)
7. SER entries: confirm each is correctly handled{drawing_instr}

IMPORTANT OUTPUT FORMAT — return ONLY this JSON, no other text:
{{
  "issues": [
    {{
      "no": <integer segment number from the source, or sequential if unavailable>,
      "original_jp": "<exact source Japanese sentence>",
      "existing_en": "<exact current English translation>",
      "area": "<明細書|請求項|要約書|図面>",
      "issue": "<comprehensive issue description in Korean — include suggested correction inline>",
      "corrected_en": "<corrected English translation, or empty string if no change needed>",
      "severity": "<심각|보통|경미>"
    }}
  ],
  "source_errors": [
    {{
      "no": <integer>,
      "original_jp": "<source text with error>",
      "area": "<area>",
      "issue": "<error description in Korean>",
      "ser_required": <true|false>,
      "ser_note": "<SER entry text in Korean, or empty string>"
    }}
  ],
  "drawing_mismatches": [
    {{
      "location": "<figure/paragraph>",
      "jp_ref": "<JP numeral>",
      "en_ref": "<EN numeral>",
      "drawing_ref": "<drawing numeral>",
      "issue": "<mismatch description in Korean>"
    }}
  ],
  "ser_verification": [
    {{
      "ser_no": "<SER item number>",
      "location": "<location>",
      "status": "<OK|미반영|추가필요>",
      "note": "<review note in Korean>"
    }}
  ],
  "summary": "<2-3 sentence overall summary in Korean>"
}}

Omit sentences/paragraphs with no issues.

[原文 (日本語)]
{source_text[:6000]}

[訳文 (English)]
{translation_text[:6000]}{notes_block}{drawing_block}{ser_block}"""

    raw = translator._call_api(
        system=[{
            "type": "text",
            "text": (
                "You are a senior patent translation verifier with deep expertise in "
                "Japanese PCT applications and US/EP patent drafting conventions. "
                "Respond only with the requested JSON object."
            ),
            "cache_control": {"type": "ephemeral"},
        }],
        messages=[{"role": "user", "content": prompt}],
    )

    try:
        match = _re.search(r'\{.*\}', raw, _re.DOTALL)
        result = json.loads(match.group() if match else raw)
        return jsonify(result)
    except Exception:
        return jsonify({"error": f"응답 파싱 오류: {raw[:300]}"}), 500


@app.route("/api/patent-report", methods=["POST"])
def api_patent_report():
    """Generate reports (DOCX + XLSX) from patent-verify results."""
    data = request.get_json(force=True)
    fmt = data.get("format", "xlsx")   # "xlsx" or "docx"
    issues = data.get("issues", [])
    source_errors = data.get("source_errors", [])
    drawing_mismatches = data.get("drawing_mismatches", [])
    ser_verification = data.get("ser_verification", [])
    summary = data.get("summary", "")
    filenames = data.get("filenames", {})

    if fmt == "xlsx":
        return _make_xlsx_report(issues, source_errors, drawing_mismatches, ser_verification, summary, filenames)
    return _make_docx_report(issues, source_errors, drawing_mismatches, ser_verification, summary, filenames)


def _make_xlsx_report(issues, source_errors, drawing_mismatches, ser_verification, summary, filenames):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    # ── 検証結果 sheet ──
    ws = wb.active
    ws.title = "検証結果"

    HDR_FILL   = PatternFill("solid", fgColor="1A73E8")
    HDR_FONT   = Font(bold=True, color="FFFFFF", size=10)
    THIN       = Side(style="thin", color="DADCE0")
    BORDER     = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    WRAP       = Alignment(wrap_text=True, vertical="top")
    SEV_FILLS  = {
        "심각": PatternFill("solid", fgColor="FCE8E6"),
        "보통": PatternFill("solid", fgColor="FEF7E0"),
        "경미": PatternFill("solid", fgColor="E6F4EA"),
    }

    headers = ["No", "原文", "訳文", "エリア", "指摘事項", "修正訳文", "심각도", "確認済"]
    col_widths = [6, 40, 40, 10, 60, 40, 8, 8]

    for c, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(1, c, h)
        cell.font = HDR_FONT
        cell.fill = HDR_FILL
        cell.alignment = WRAP
        cell.border = BORDER
        ws.column_dimensions[get_column_letter(c)].width = w

    ws.row_dimensions[1].height = 20

    for i, item in enumerate(issues, 2):
        sev = item.get("severity", "경미")
        vals = [
            item.get("no", i - 1),
            item.get("original_jp", ""),
            item.get("existing_en", ""),
            item.get("area", ""),
            item.get("issue", ""),
            item.get("corrected_en", ""),
            sev,
            "",
        ]
        fill = SEV_FILLS.get(sev)
        for c, v in enumerate(vals, 1):
            cell = ws.cell(i, c, v)
            cell.alignment = WRAP
            cell.border = BORDER
            if fill and c in (1, 4, 7, 8):
                cell.fill = fill

    # ── 原文エラー sheet ──
    if source_errors:
        ws2 = wb.create_sheet("原文エラー")
        hdrs2 = ["No", "原文", "エリア", "エラー内容", "SER必要", "SER記載"]
        widths2 = [6, 40, 10, 60, 8, 50]
        for c, (h, w) in enumerate(zip(hdrs2, widths2), 1):
            cell = ws2.cell(1, c, h)
            cell.font = HDR_FONT
            cell.fill = PatternFill("solid", fgColor="EA4335")
            cell.alignment = WRAP; cell.border = BORDER
            ws2.column_dimensions[get_column_letter(c)].width = w
        for i, e in enumerate(source_errors, 2):
            for c, v in enumerate([
                e.get("no", i-1), e.get("original_jp",""), e.get("area",""),
                e.get("issue",""), "✓" if e.get("ser_required") else "", e.get("ser_note","")
            ], 1):
                cell = ws2.cell(i, c, v); cell.alignment = WRAP; cell.border = BORDER

    # ── SER検証 sheet ──
    if ser_verification:
        ws3 = wb.create_sheet("SER検証")
        hdrs3 = ["SER No.", "位置", "ステータス", "検討内容"]
        widths3 = [10, 20, 12, 70]
        for c, (h, w) in enumerate(zip(hdrs3, widths3), 1):
            cell = ws3.cell(1, c, h)
            cell.font = HDR_FONT
            cell.fill = PatternFill("solid", fgColor="34A853")
            cell.alignment = WRAP; cell.border = BORDER
            ws3.column_dimensions[get_column_letter(c)].width = w
        for i, s in enumerate(ser_verification, 2):
            for c, v in enumerate([s.get("ser_no",""), s.get("location",""), s.get("status",""), s.get("note","")], 1):
                cell = ws3.cell(i, c, v); cell.alignment = WRAP; cell.border = BORDER

    # ── 図面不一致 sheet ──
    if drawing_mismatches:
        ws4 = wb.create_sheet("図面不一致")
        hdrs4 = ["位置", "JP符号", "EN符号", "図面符号", "内容"]
        widths4 = [15, 12, 12, 12, 60]
        for c, (h, w) in enumerate(zip(hdrs4, widths4), 1):
            cell = ws4.cell(1, c, h)
            cell.font = HDR_FONT
            cell.fill = PatternFill("solid", fgColor="F9AB00")
            cell.alignment = WRAP; cell.border = BORDER
            ws4.column_dimensions[get_column_letter(c)].width = w
        for i, m in enumerate(drawing_mismatches, 2):
            for c, v in enumerate([m.get("location",""), m.get("jp_ref",""), m.get("en_ref",""), m.get("drawing_ref",""), m.get("issue","")], 1):
                cell = ws4.cell(i, c, v); cell.alignment = WRAP; cell.border = BORDER

    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    return send_file(buf,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True, download_name="번역검증보고서.xlsx")


def _make_docx_report(issues, source_errors, drawing_mismatches, ser_verification, summary, filenames):
    import docx
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    doc.add_heading("特許翻訳検証報告書 / 특허 번역 검증 보고서", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if filenames:
        for k, v in filenames.items():
            if v: doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("検証要約 / 검증 요약", 1)
    doc.add_paragraph(summary or "—")

    def _hdr_cell(cell, text, rgb_hex):
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn; from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"), rgb_hex)
        cell._tc.get_or_add_tcPr().append(shd)

    doc.add_heading(f"指摘事項一覧 / 지적사항 ({len(issues)}건)", 1)
    if issues:
        tbl = doc.add_table(rows=1, cols=7); tbl.style = "Table Grid"
        for i, h in enumerate(["No","原文","訳文","エリア","指摘事項","修正訳文","심각도"]):
            _hdr_cell(tbl.rows[0].cells[i], h, "1A73E8")
        for item in issues:
            r = tbl.add_row().cells
            r[0].text = str(item.get("no",""))
            r[1].text = item.get("original_jp","")
            r[2].text = item.get("existing_en","")
            r[3].text = item.get("area","")
            r[4].text = item.get("issue","")
            r[5].text = item.get("corrected_en","")
            r[6].text = item.get("severity","")
    else:
        doc.add_paragraph("지적사항 없음")

    doc.add_heading(f"原文エラー / 원문 오류 ({len(source_errors)}건)", 1)
    if source_errors:
        tbl = doc.add_table(rows=1, cols=5); tbl.style = "Table Grid"
        for i, h in enumerate(["No","原文","エリア","エラー内容","SER"]):
            _hdr_cell(tbl.rows[0].cells[i], h, "EA4335")
        for e in source_errors:
            r = tbl.add_row().cells
            r[0].text = str(e.get("no","")); r[1].text = e.get("original_jp","")
            r[2].text = e.get("area",""); r[3].text = e.get("issue","")
            r[4].text = "필요" if e.get("ser_required") else "참고"
    else:
        doc.add_paragraph("원문 오류 없음")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return send_file(buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True, download_name="번역검증보고서.docx")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=True, host="0.0.0.0", port=port)
